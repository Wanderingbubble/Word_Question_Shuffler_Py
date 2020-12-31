import docx
import re
import random
from random import randint

file = ".\InputWordFiles\sample.docx"                   #[Enter the file path to the word file you want to process]

def getText(filename):
    orginaldoc = docx.Document(filename)
    fullArray = []                           #created array
    for i in orginaldoc.paragraphs:
        fullArray.append(i.text)
    return "\n".join(fullArray)              #takes each para from array and combined "\n".join

def random_with_N_digits(n):
    range_start = 10**(n-1)
    range_end = (10**n)-1
    return randint(range_start, range_end)

def makeFile(createdarray):                  #Function which rejoins the para with the help of index numbers and saves the file as Randomized.docx
    newDoc = docx.Document()
    for i in createdarray:
        newDoc.add_paragraph(i)
        newDoc.add_page_break()
    newDoc.save(f'.\ShuffledWordFiles\Randomized{random_with_N_digits(4)}.docx')

Array = getText(file)                              
print(type(Array))                           #to debug
WQarray = re.split(r'\d\d\)\-\-', Array)     #Splits the paras list
Orignallist = WQarray                        #saves the orignal list if needed for further process
random.shuffle(WQarray)                      #Uses python's random module to shuffle the list
makeFile(WQarray)                            #Passes the list in makeFile() which writes a new word document named randomized.docx   


#OLD CODE----------
# print(WQarray)
# EWQarray = list(enumerate(WQarray, start=1))
# for i in WQarray:
    # print(i)

# regxx = re.compile(r'\d\d\)')
# newlist = list(filter(regxx.match, Array))
# print(newlist)

# def ListSpitter(listName):
#     splitArray = []
#     insideList = []
#     for i in Array:
#         insideList.append(i)
#         if i = filter(regxx.match, Array)
# print(Array)
# TotalText = getText(p2)

# def demoTextToDoc(textName):
#     demodoc = docx.Document(textName)
#     demodoc.save(tex.docx") #Creates a file with textName vairable

# demoTextToDoc(TotalText)
